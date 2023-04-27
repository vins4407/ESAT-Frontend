import streamlit as st
from io import BytesIO
from docx import Document
from helper import defaultConfig
from docx.shared import Inches
import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
import plotly.express as px 
import plotly.graph_objs as go
import plotly.io as pio
import pandas as pd
from services import get_userdeatils
import random


defaultConfig()

def Report():
    st.title("Generate Report :memo:")
    
    st.selectbox("Select the IP",["192.168.211.136"])
    ip ="192.168.211.136"
    mac="94:08:53:9b:37:d7"

    vulnerabilities = [
    "SQL injection",
    "Cross-site scripting (XSS)",
    "Cross-site request forgery (CSRF)",
    "Remote code execution (RCE)",
    "Denial of service (DoS)",
    "Man-in-the-middle (MitM)",
    "Clickjacking",
    "File inclusion",
    "Buffer overflow",
    "Authentication bypass"
]


    domains = ["example.com", "google.com", "facebook.com", "stackoverflow.com"]
    cves = [
    "CVE-2017-5638",
    "CVE-2017-11774",
    "CVE-2019-11510",
    "CVE-2019-19781",
    "CVE-2020-1472",
    "CVE-2021-26855",
    "CVE-2021-21985",
    "CVE-2021-34527",
    "CVE-2021-33742",
    "CVE-2021-40444"
]
    df = pd.DataFrame(
        [
        {"command": "nmap", "rating": 4, "Select": True},
        {"command": "tshark", "rating": 5, "Select": False},
        {"command": "openvas", "rating": 3, "Select": True},
    ]
        )
    ip ="10.76.15.153"
    mac="94:08:53:9b:37:d7"
    n=2
    
    open_ports = px.bar(data_frame=df,x="command",y="rating")

    document = Document()
    
    document.add_heading('E-sat Report', 0)
    para = document.add_paragraph('The EMPLOYEE SECURITY AWARENESS TESTING TOOL is a security testing tool aimed at testing the security awareness of employees in an organization. The tool is designed to run on a local network and tests the security of the employees devices connected to the network, without their knowl')
    para.paragraph_format.line_spacing = Inches(0.2)
    img_bytes = pio.to_image(open_ports, format='png')
    stream = BytesIO(img_bytes)

    document.add_picture(stream, width=Inches(6))

    
     # Add current date and time in the header
    header = document.sections[0].header
    paragraph = header.paragraphs[0]
    paragraph.text = datetime.datetime.now().strftime("%B %d, %Y %H:%M:%S")
    paragraph.alignment = 2
    
    document.add_heading(f'IP :- {ip}', 3,)
    document.add_heading('MAC :- 94:08:53:9b:37:d7', 3)
    
    document.add_heading('RATING : 4', 3)
    document.add_heading('VULNERABILITY', 3,)
    table = document.add_table(rows=n, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set the column names
    table.rows[0].cells[0].text = "RISK LEVEL "
    table.rows[0].cells[1].text = "VULNERABILITY NAME"
    table.rows[0].cells[2].text = "CV"

   
    # Add some sample data to the table
    for i in range(1,n):
        table.rows[i].cells[0].text = f"{int(random.uniform(2, 5))}"
        table.rows[i].cells[1].text = f"{random.choice(vulnerabilities)}"
        table.rows[i].cells[2].text = f"{random.choice(cves)}"
    for row in table.rows:
        for cell in row.cells:
            cell.margin_left = Cm(0.25)
            cell.margin_right = Cm(0.25)
    
    #domain
    document.add_heading('TRAFFIC', 3)
    
    for domain in domains:
        p = document.add_paragraph(domain)
    # Center the text within the paragraph
    

    document.add_heading('suspicious domains', 3)
    
    # Add a table with n rows and 2 columns1
    table = document.add_table(rows=n, cols=2)

    # Set the column names
    table.rows[0].cells[0].text = "SOURCES"
    table.rows[0].cells[1].text = "STATUS"
    for row in table.rows:
        for cell in row.cells:
            cell.margin_left = Cm(0.50)
            cell.margin_right = Cm(0.50)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


    # Add some sample data to the table
    for i in range(1,n):
        table.rows[i].cells[0].text = f"ecil.com{i}"
        table.rows[i].cells[1].text = "Unrated"
    
    
    document.add_heading('PORTS', 3)
    table = document.add_table(rows=n, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set the column names
    table.rows[0].cells[0].text = "NUMBER"
    table.rows[0].cells[1].text = "SERVICES"
    table.rows[0].cells[2].text = "INFO"

    # Add some sample data to the table
    for i in range(1,n):
        table.rows[i].cells[0].text = f"8021{i}"
        table.rows[i].cells[1].text = "http"
        table.rows[i].cells[2].text = "alternat https proxy port"
    
    stream = BytesIO()

    document.save(stream)
    stream.seek(0)
                
    st.download_button(
                    label="Download report file",
                    data= stream,
                    file_name='report.docx',
                    mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )


if __name__ == "__main__":
    if st.session_state["authentication_status"]:
        Report()
    else:
        st.warning("Logging to acess the page")